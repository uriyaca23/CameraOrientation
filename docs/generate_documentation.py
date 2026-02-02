"""
Complete Documentation Generator with OMML Equations Inline.
This is a unified script that generates the FULL original document with proper Word equations.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


# =============================================================================
# OMML Equation Helper Functions
# =============================================================================

def add_math_text(omath, text):
    """Add text to an oMath element."""
    run = OxmlElement('m:r')
    t = OxmlElement('m:t')
    t.text = text
    run.append(t)
    omath.append(run)


def create_inline_math(text):
    """Create an inline oMath element."""
    omath = OxmlElement('m:oMath')
    add_math_text(omath, text)
    return omath


def add_para_with_inline_eq(doc, parts):
    """
    Add a paragraph with mixed text and inline equations.
    parts is a list of tuples: ('text', 'some text') or ('eq', 'equation')
    """
    para = doc.add_paragraph()
    for part_type, content in parts:
        if part_type == 'text':
            para.add_run(content)
        elif part_type == 'eq':
            omath = create_inline_math(content)
            para._p.append(omath)
    return para


def add_display_equation(doc, eq_text, style_after=None):
    """Add a centered display equation."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omathpara = OxmlElement('m:oMathPara')
    omath = OxmlElement('m:oMath')
    add_math_text(omath, eq_text)
    omathpara.append(omath)
    para._p.append(omathpara)
    return para


def add_bullet_with_eq(doc, parts):
    """Add a bullet point with inline equations."""
    para = add_para_with_inline_eq(doc, parts)
    para.style = 'List Bullet'
    return para


# =============================================================================
# CREATE DOCUMENT
# =============================================================================
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('CameraOrientation: Smartphone IMU-Based Orientation Estimation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('A Comprehensive Technical Documentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.italic = True

doc.add_paragraph()
author = doc.add_paragraph('Project Documentation')
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
date = doc.add_paragraph('February 2026')
date.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Executive Summary",
    "2. Introduction and Motivation",
    "   2.1 The Orientation Estimation Problem",
    "   2.2 Sensor Fusion Paradigm",
    "   2.3 Project Objectives",
    "   2.4 Applications and Use Cases",
    "3. Theoretical Background",
    "   3.1 Orientation Representations",
    "   3.2 IMU Sensor Physics",
    "   3.3 Sensor Noise Characterization",
    "4. Mathematical Framework",
    "   4.1 Extended Kalman Filter (EKF)",
    "   4.2 Factor Graph Optimization",
    "   4.3 Quaternion Kinematics",
    "5. Implementation Details",
    "   5.1 Google EKF Solver",
    "   5.2 PyTorch Factor Graph Solver",
    "   5.3 Noise Database Architecture",
    "6. Research Process and Challenges",
    "   6.1 Synchronization Problem",
    "   6.2 Bias Estimation Difficulties",
    "   6.3 Solver Comparison Analysis",
    "7. Results and Analysis",
    "   7.1 Synthetic Data Validation",
    "   7.2 Real-World Performance",
    "   7.3 Failure Mode Analysis",
    "8. Usage Guide",
    "   8.1 Installation",
    "   8.2 Quick Start",
    "   8.3 API Reference",
    "   8.4 Configuration Options",
    "9. Project Architecture",
    "10. References",
    "Appendix A: Complete Noise Database",
    "Appendix B: Code Examples",
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# ============================================================================
# CHAPTER 1: EXECUTIVE SUMMARY
# ============================================================================
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'This document provides comprehensive technical documentation for the CameraOrientation project, '
    'a sophisticated software system designed to estimate the 3D orientation of smartphones using '
    'onboard Inertial Measurement Unit (IMU) sensors. The project addresses one of the fundamental '
    'challenges in mobile robotics, augmented reality, and video stabilization: accurately determining '
    'the spatial orientation of a device in real-time using noisy, biased sensor measurements.'
)

doc.add_paragraph(
    'The system implements two distinct algorithmic approaches to orientation estimation: an Extended '
    'Kalman Filter (EKF) based on the Android Open Source Project (AOSP) sensor fusion implementation, '
    'and a Factor Graph optimization approach using PyTorch for automatic differentiation. Both methods '
    'fuse data from three complementary sensors: the triaxial gyroscope (angular velocity), triaxial '
    'accelerometer (specific force/gravity), and triaxial magnetometer (Earth\'s magnetic field).'
)

doc.add_paragraph('The project delivers several key contributions:')

contributions = [
    'A comprehensive noise parameter database covering 100+ smartphone models with datasheet-derived specifications',
    'Two production-ready orientation estimation algorithms with uncertainty quantification',
    'A simple, easy-to-use Python API for IMU-to-orientation conversion',
    'Detailed analysis tools for sensor synchronization, debugging, and visualization',
    'Video generation capabilities with overlaid orientation visualization',
]
for c in contributions:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph(
    'Performance validation on real-world data from a Google Pixel 10 smartphone demonstrates that '
    'the Google EKF implementation achieves orientation accuracy within 2-5 degrees under typical '
    'indoor conditions, with the factor graph approach providing smoother trajectories at the cost '
    'of increased computational requirements. The system successfully handles challenging scenarios '
    'including rapid rotations, magnetic disturbances, and sensor bias drift.'
)

doc.add_paragraph(
    'This documentation covers the complete theoretical foundation, mathematical derivations, '
    'implementation details, research process, experimental results, and practical usage guidelines. '
    'It is intended for researchers, engineers, and developers working in sensor fusion, robotics, '
    'augmented reality, or related fields who require a deep understanding of smartphone-based '
    'orientation estimation.'
)

doc.add_page_break()

# ============================================================================
# CHAPTER 2: INTRODUCTION AND MOTIVATION
# ============================================================================
doc.add_heading('2. Introduction and Motivation', level=1)

doc.add_heading('2.1 The Orientation Estimation Problem', level=2)

doc.add_paragraph(
    'Determining the orientation of a rigid body in three-dimensional space is a fundamental problem '
    'that arises in numerous applications across robotics, aerospace, virtual reality, augmented reality, '
    'video stabilization, and mobile computing. In the context of smartphones, orientation estimation '
    'enables features such as automatic screen rotation, compass applications, panoramic photography, '
    'motion-controlled gaming, and inertial navigation.'
)

doc.add_paragraph(
    'The orientation of a smartphone can be described as the rotational relationship between two '
    'coordinate frames: the body frame (fixed to the device) and the world frame (typically aligned '
    'with Earth\'s local East-North-Up or North-East-Down reference). This relationship is commonly '
    'expressed using one of several mathematical representations including Euler angles (yaw, pitch, roll), '
    'rotation matrices, quaternions, or rotation vectors. Each representation has distinct advantages '
    'and limitations that make it suitable for different aspects of the estimation problem.'
)

doc.add_paragraph(
    'The fundamental challenge in orientation estimation arises from the complementary strengths and '
    'weaknesses of the available sensors:'
)

# Sensor comparison table
sensors_table = doc.add_table(rows=4, cols=3)
sensors_table.style = 'Table Grid'
hdr_cells = sensors_table.rows[0].cells
hdr_cells[0].text = 'Sensor'
hdr_cells[1].text = 'Strengths'
hdr_cells[2].text = 'Weaknesses'

row1 = sensors_table.rows[1].cells
row1[0].text = 'Gyroscope'
row1[1].text = 'High bandwidth, measures rotation directly, immune to external forces'
row1[2].text = 'Integrates to position (drift), bias instability, no absolute reference'

row2 = sensors_table.rows[2].cells
row2[0].text = 'Accelerometer'
row2[1].text = 'Absolute gravity reference, no drift, constrains pitch/roll'
row2[2].text = 'Cannot distinguish gravity from acceleration, sensitive to vibration'

row3 = sensors_table.rows[3].cells
row3[0].text = 'Magnetometer'
row3[1].text = 'Absolute heading reference (magnetic north), constrains yaw'
row3[2].text = 'Easily disturbed by ferromagnetic materials, indoor interference'

doc.add_paragraph()

doc.add_heading('2.2 Sensor Fusion Paradigm', level=2)

doc.add_paragraph(
    'The solution to the orientation estimation problem lies in sensor fusion - the intelligent '
    'combination of multiple sensor measurements to produce an estimate that is more accurate and '
    'robust than any single sensor could provide alone. The key insight is that the sensors provide '
    'complementary information:'
)

doc.add_paragraph(
    'The gyroscope provides high-frequency, low-latency measurements of angular velocity that can be '
    'integrated to track orientation changes over short time periods. However, any small bias in the '
    'gyroscope reading will cause the integrated orientation to drift unboundedly over time. A bias '
    'of just 0.1 degrees per second will accumulate to 6 degrees of error after one minute and 360 '
    'degrees (a full rotation) after one hour.'
)

doc.add_paragraph(
    'The accelerometer measures the specific force acting on the sensor, which in the absence of '
    'linear acceleration equals the gravitational acceleration vector pointing "up" in the world frame. '
    'By observing the direction of gravity in the body frame, we can determine two of the three degrees '
    'of freedom in orientation (pitch and roll). However, the accelerometer cannot distinguish between '
    'gravitational and inertial acceleration, and provides no information about rotation around the '
    'vertical axis (yaw/heading).'
)

doc.add_paragraph(
    'The magnetometer measures the Earth\'s magnetic field vector, which points toward magnetic north '
    'with a downward inclination (in the northern hemisphere). By observing the magnetic field direction '
    'in the body frame and combining it with the accelerometer-derived pitch/roll, we can determine the '
    'absolute heading. However, magnetometers are highly susceptible to interference from ferromagnetic '
    'materials, electrical currents, and magnetic anomalies that are pervasive in indoor environments.'
)

doc.add_heading('2.3 Project Objectives', level=2)

doc.add_paragraph('This project was initiated with the following primary objectives:')

objectives = [
    'Implement a robust, production-quality orientation estimation system for smartphone IMU data',
    'Compare and contrast recursive filtering (EKF) and batch optimization (factor graph) approaches',
    'Create a comprehensive noise parameter database for accurate sensor modeling across devices',
    'Provide uncertainty quantification (covariance estimates) alongside orientation estimates',
    'Generate high-quality visualizations for debugging, analysis, and demonstration purposes',
    'Document the complete theoretical foundation and implementation details for reproducibility',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_heading('2.4 Applications and Use Cases', level=2)

doc.add_paragraph('The orientation estimation capabilities developed in this project have broad applicability:')

doc.add_paragraph(
    'Video Stabilization: By knowing the camera\'s orientation during recording, we can apply '
    'compensating rotations to remove unwanted camera shake and produce smooth, professional-looking '
    'video. The orientation covariance estimates enable adaptive stabilization that preserves '
    'intentional camera movements while filtering out high-frequency jitter.'
)

doc.add_paragraph(
    'Augmented Reality: AR applications require precise knowledge of the device\'s orientation to '
    'correctly render virtual objects that appear fixed in the real world. The low-latency orientation '
    'estimates from the EKF are particularly suitable for this application, while the uncertainty '
    'estimates can drive visual feedback about tracking quality.'
)

doc.add_paragraph(
    'Inertial Navigation: When GPS signals are unavailable (indoors, tunnels, urban canyons), '
    'smartphone orientation combined with step detection can provide pedestrian dead reckoning '
    'capabilities. Accurate orientation is essential for correctly projecting measured accelerations '
    'into the navigation frame.'
)

doc.add_paragraph(
    'Motion Capture: For applications ranging from sports analysis to rehabilitation monitoring, '
    'smartphone orientation can serve as a low-cost alternative to dedicated motion capture systems. '
    'The ability to record and replay orientation trajectories enables detailed biomechanical analysis.'
)

doc.add_page_break()

# Save Part 1 and print progress
# (save moved to end)
# Part 1 content added


# ============================================================================
# CHAPTER 3: THEORETICAL BACKGROUND
# ============================================================================
doc.add_heading('3. Theoretical Background', level=1)

doc.add_paragraph(
    'Before diving into the estimation algorithms, it is essential to understand the mathematical '
    'representations used to describe orientation in three-dimensional space. Each representation '
    'has distinct properties that make it suitable for different purposes.'
)

doc.add_heading('3.1 Orientation Representations', level=2)

doc.add_heading('3.1.1 Rotation Matrices', level=3)

add_para_with_inline_eq(doc, [
    ('text', 'A rotation matrix '),
    ('eq', 'R ∈ SO(3)'),
    ('text', ' is a 3×3 orthogonal matrix with determinant +1 that transforms '
             'vectors from one coordinate frame to another. The set of all such matrices forms the Special '
             'Orthogonal group SO(3). Key properties include:')
])

# Properties with inline equations
add_bullet_with_eq(doc, [('eq', 'RᵀR = RRᵀ = I'), ('text', ' (orthogonality)')])
add_bullet_with_eq(doc, [('eq', 'det(R) = +1'), ('text', ' (proper rotation, not reflection)')])
add_bullet_with_eq(doc, [('eq', 'R⁻¹ = Rᵀ'), ('text', ' (inverse equals transpose)')])
add_bullet_with_eq(doc, [('text', 'Composition: '), ('eq', 'R_AC = R_AB · R_BC')])

add_para_with_inline_eq(doc, [
    ('text', 'The rotation matrix from world frame to body frame can be written as '),
    ('eq', 'R_wb'),
    ('text', ', where a vector expressed in the world frame '),
    ('eq', 'v_w'),
    ('text', ' is transformed to the body frame as '),
    ('eq', 'v_b = R_wb · v_w'),
    ('text', '. The columns of '),
    ('eq', 'R_wb'),
    ('text', ' represent the world frame basis vectors expressed in body frame coordinates.')
])

doc.add_paragraph(
    'While rotation matrices are intuitive and compose naturally through matrix multiplication, '
    'they have significant drawbacks for estimation: they require 9 parameters with 6 constraints '
    '(orthogonality), making them over-parameterized and requiring constrained optimization or '
    'projection steps to maintain validity.'
)

doc.add_heading('3.1.2 Euler Angles (Yaw-Pitch-Roll)', level=3)

doc.add_paragraph(
    'Euler angles represent orientation as a sequence of three rotations about specified axes. '
    'The most common convention for aerospace and robotics applications is the ZYX (yaw-pitch-roll) '
    'sequence, also known as Tait-Bryan angles:'
)

add_bullet_with_eq(doc, [('text', 'Yaw ('), ('eq', 'ψ'), ('text', '): Rotation about the vertical (Z) axis, representing heading/azimuth')])
add_bullet_with_eq(doc, [('text', 'Pitch ('), ('eq', 'θ'), ('text', '): Rotation about the lateral (Y) axis, representing nose-up/down attitude')])
add_bullet_with_eq(doc, [('text', 'Roll ('), ('eq', 'φ'), ('text', '): Rotation about the longitudinal (X) axis, representing bank angle')])

doc.add_paragraph('The rotation matrix corresponding to ZYX Euler angles is:')

add_display_equation(doc, 'R = R_z(ψ) · R_y(θ) · R_x(φ)')

add_para_with_inline_eq(doc, [
    ('text', 'Euler angles are intuitive for human interpretation and require only 3 parameters, but suffer '
             'from gimbal lock - a singularity that occurs when '),
    ('eq', 'θ = ±90°'),
    ('text', ', where yaw and roll become indistinguishable. This makes Euler angles unsuitable for '
             'estimation algorithms that must handle arbitrary orientations. They are used only for output/display purposes in this project.')
])

doc.add_heading('3.1.3 Unit Quaternions', level=3)

add_para_with_inline_eq(doc, [
    ('text', 'Unit quaternions provide a singularity-free, computationally efficient representation of '
             'orientation that is ideal for estimation algorithms. A quaternion '),
    ('eq', 'q = [w, x, y, z]ᵀ'),
    ('text', ' consists of a scalar part '),
    ('eq', 'w'),
    ('text', ' and a vector part '),
    ('eq', 'v = [x, y, z]ᵀ'),
    ('text', ', subject to the unit constraint:')
])

add_display_equation(doc, '‖q‖² = w² + x² + y² + z² = 1')

add_para_with_inline_eq(doc, [
    ('text', 'A unit quaternion can be interpreted geometrically as encoding a rotation of angle '),
    ('eq', 'θ'),
    ('text', ' about a unit axis '),
    ('eq', 'n'),
    ('text', ':')
])

add_display_equation(doc, 'q = [cos(θ/2), sin(θ/2)·nₓ, sin(θ/2)·nᵧ, sin(θ/2)·n_z]ᵀ')

doc.add_paragraph('Key quaternion operations used in this project include:')

doc.add_paragraph(
    'Quaternion Multiplication: The composition of two rotations is computed as quaternion '
    'multiplication using the Hamilton product:'
)

add_display_equation(doc, 'q₁ ⊗ q₂ = [w₁w₂ − v₁·v₂, w₁v₂ + w₂v₁ + v₁×v₂]ᵀ')

add_para_with_inline_eq(doc, [
    ('text', 'Quaternion to Rotation Matrix: The rotation matrix corresponding to quaternion '),
    ('eq', 'q'),
    ('text', ' is:')
])

add_display_equation(doc, 'R(q) = (w² − ‖v‖²)I + 2vvᵀ + 2w[v]ₓ')

add_para_with_inline_eq(doc, [
    ('text', 'where '),
    ('eq', '[v]ₓ'),
    ('text', ' is the skew-symmetric matrix of '),
    ('eq', 'v'),
    ('text', '. Note that '),
    ('eq', 'q'),
    ('text', ' and '),
    ('eq', '−q'),
    ('text', ' represent the same rotation, so a sign convention (typically '),
    ('eq', 'w > 0'),
    ('text', ') is often enforced for uniqueness.')
])

doc.add_heading('3.2 IMU Sensor Physics', level=2)

doc.add_heading('3.2.1 Gyroscope Model', level=3)

doc.add_paragraph(
    'A MEMS (Micro-Electro-Mechanical Systems) gyroscope measures angular velocity by detecting '
    'the Coriolis force acting on a vibrating proof mass. The measurement model for a triaxial '
    'gyroscope is:'
)

add_display_equation(doc, 'ω_meas = ω_true + b_g + n_g')

doc.add_paragraph('where:')

add_bullet_with_eq(doc, [('eq', 'ω_meas ∈ ℝ³'), ('text', ' is the measured angular velocity [rad/s]')])
add_bullet_with_eq(doc, [('eq', 'ω_true ∈ ℝ³'), ('text', ' is the true angular velocity of the body frame relative to the world frame')])
add_bullet_with_eq(doc, [('eq', 'b_g ∈ ℝ³'), ('text', ' is the gyroscope bias - a slowly-varying offset that must be estimated')])
add_bullet_with_eq(doc, [('eq', 'n_g ~ N(0, σ_g²I)'), ('text', ' is white Gaussian measurement noise')])

doc.add_paragraph(
    'The gyroscope bias is not constant but drifts slowly over time due to temperature changes '
    'and other environmental factors. This drift is typically modeled as a random walk:'
)

add_display_equation(doc, 'ḃ_g = n_bg, where n_bg ~ N(0, σ_bg²I)')

add_para_with_inline_eq(doc, [
    ('text', 'The bias drift rate '),
    ('eq', 'σ_bg'),
    ('text', ' is typically orders of magnitude smaller than the measurement noise '),
    ('eq', 'σ_g'),
    ('text', ', meaning the bias changes slowly compared to the measurement rate. For consumer-grade MEMS '
             'gyroscopes, typical values are '),
    ('eq', 'σ_g ≈ 0.001-0.01 rad/s'),
    ('text', ' and '),
    ('eq', 'σ_bg ≈ 10⁻⁵-10⁻⁴ rad/s/√s'),
    ('text', '.')
])

doc.add_heading('3.2.2 Accelerometer Model', level=3)

doc.add_paragraph(
    'A MEMS accelerometer measures specific force - the non-gravitational force per unit mass '
    'acting on the sensor. When the device is stationary or moving at constant velocity, the '
    'only force is the reaction to gravity. The measurement model is:'
)

add_display_equation(doc, 'a_meas = R_wb · (a_true − g) + b_a + n_a')

doc.add_paragraph('where:')

add_bullet_with_eq(doc, [('eq', 'a_meas ∈ ℝ³'), ('text', ' is the measured specific force in body frame [m/s²]')])
add_bullet_with_eq(doc, [('eq', 'R_wb'), ('text', ' is the rotation from world to body frame')])
add_bullet_with_eq(doc, [('eq', 'a_true'), ('text', ' is the true linear acceleration in world frame')])
add_bullet_with_eq(doc, [('eq', 'g = [0, 0, −9.81]ᵀ'), ('text', ' is gravity in world frame (pointing down)')])
add_bullet_with_eq(doc, [('eq', 'b_a ∈ ℝ³'), ('text', ' is the accelerometer bias')])
add_bullet_with_eq(doc, [('eq', 'n_a ~ N(0, σ_a²I)'), ('text', ' is white Gaussian measurement noise')])

add_para_with_inline_eq(doc, [
    ('text', 'When the device is stationary ('),
    ('eq', 'a_true = 0'),
    ('text', '), the accelerometer measures the reaction force to gravity, which points upward in the body frame. '
             'This provides an absolute reference for the pitch and roll angles. However, during dynamic motion '
             '(walking, driving), the accelerometer reads a combination of gravity and linear acceleration that '
             'cannot be separated without additional information.')
])

doc.add_heading('3.2.3 Magnetometer Model', level=3)

doc.add_paragraph(
    'The magnetometer measures the Earth\'s magnetic field vector in the body frame. In the '
    'local world frame (ENU - East-North-Up), the magnetic field has a horizontal component '
    'pointing toward magnetic north and a vertical component whose direction depends on the '
    'magnetic inclination. The measurement model is:'
)

add_display_equation(doc, 'm_meas = R_wb · m_earth + b_m + n_m')

doc.add_paragraph('where:')

add_bullet_with_eq(doc, [('eq', 'm_meas ∈ ℝ³'), ('text', ' is the measured magnetic field in body frame [µT]')])
add_bullet_with_eq(doc, [('eq', 'm_earth'), ('text', ' is the Earth\'s magnetic field in world frame')])
add_bullet_with_eq(doc, [('eq', 'b_m ∈ ℝ³'), ('text', ' is hard-iron bias from permanent magnets near the sensor')])
add_bullet_with_eq(doc, [('eq', 'n_m ~ N(0, σ_m²I)'), ('text', ' is measurement noise (includes soft-iron effects as approximation)')])

add_para_with_inline_eq(doc, [
    ('text', 'The magnetometer is the most challenging sensor to use reliably. In indoor environments, '
             'ferromagnetic structural elements (rebar, HVAC ducts, furniture) create local magnetic '
             'anomalies that can cause errors of tens of degrees in heading. The noise model used in '
             'this project assigns much higher uncertainty to magnetometer measurements indoors ('),
    ('eq', 'σ_m ≈ 50 µT'),
    ('text', ') compared to outdoors ('),
    ('eq', 'σ_m ≈ 5 µT'),
    ('text', ').')
])

doc.add_heading('3.3 Sensor Noise Characterization', level=2)

doc.add_paragraph(
    'Accurate sensor noise characterization is critical for optimal sensor fusion. The noise '
    'parameters determine the relative weighting of different sensor contributions and directly '
    'impact estimation accuracy and uncertainty quantification.'
)

doc.add_heading('3.3.1 Allan Variance Analysis', level=3)

add_para_with_inline_eq(doc, [
    ('text', 'The Allan variance (or Allan deviation) is the standard method for characterizing inertial '
             'sensor noise. It was originally developed for analyzing frequency standards (atomic clocks) '
             'and is now widely used for gyroscopes and accelerometers. The Allan variance '),
    ('eq', 'σ²(τ)'),
    ('text', ' is computed as a function of averaging time '),
    ('eq', 'τ'),
    ('text', ' and reveals different noise processes at different time scales:')
])

add_bullet_with_eq(doc, [('text', 'White Noise (Angle Random Walk for gyro, Velocity Random Walk for accel): '), ('eq', 'σ(τ) ∝ τ^(−1/2)')])
add_bullet_with_eq(doc, [('text', 'Bias Instability: '), ('eq', 'σ(τ) ≈ constant'), ('text', ' at the minimum of the Allan deviation curve')])
add_bullet_with_eq(doc, [('text', 'Random Walk: '), ('eq', 'σ(τ) ∝ τ^(+1/2)')])
add_bullet_with_eq(doc, [('text', 'Rate Ramp (deterministic drift): '), ('eq', 'σ(τ) ∝ τ')])

add_para_with_inline_eq(doc, [
    ('text', 'The noise density values in our database (e.g., 0.007 dps/√Hz for BMI270 gyroscope) correspond '
             'to the white noise coefficient extracted from Allan variance analysis. To convert noise density '
             'to measurement standard deviation at a given sampling rate '),
    ('eq', 'f'),
    ('text', ', we use: '),
    ('eq', 'σ = noise_density × √f'),
    ('text', '.')
])

doc.add_heading('3.3.2 Noise Database Design', level=3)

doc.add_paragraph(
    'The noise database implemented in this project maps smartphone models to their IMU sensor '
    'specifications. Since manufacturers rarely publish detailed noise specifications, we relied on '
    'three information sources:'
)

sources = [
    'IMU chip datasheets (Bosch BMI270, STM LSM6DSR, TDK ICM-42688, etc.)',
    'Academic papers characterizing smartphone IMU performance',
    'Empirical measurements from Allan variance analysis on specific devices',
]
for s in sources:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph('The database stores the following parameters for each sensor:')

params_table = doc.add_table(rows=8, cols=3)
params_table.style = 'Table Grid'
params_table.rows[0].cells[0].text = 'Parameter'
params_table.rows[0].cells[1].text = 'Units'
params_table.rows[0].cells[2].text = 'Description'

param_data = [
    ('gyro_noise_sigma', 'rad/s', 'Gyroscope white noise standard deviation'),
    ('accel_noise_sigma', 'm/s²', 'Accelerometer white noise standard deviation'),
    ('mag_noise_sigma', 'µT', 'Magnetometer noise (environment-dependent)'),
    ('gyro_bias_instability', 'rad/s', 'Gyroscope bias drift rate'),
    ('accel_bias_instability', 'm/s²', 'Accelerometer bias drift rate'),
    ('gyro_bias_sigma', 'rad/s', 'Prior uncertainty on gyroscope bias'),
    ('accel_bias_sigma', 'm/s²', 'Prior uncertainty on accelerometer bias'),
]

for i, (param, units, desc) in enumerate(param_data, 1):
    params_table.rows[i].cells[0].text = param
    params_table.rows[i].cells[1].text = units
    params_table.rows[i].cells[2].text = desc

doc.add_page_break()

# Save Part 2


# ============================================================================
# CHAPTER 4: MATHEMATICAL FRAMEWORK
# ============================================================================
doc.add_heading('4. Mathematical Framework', level=1)

doc.add_paragraph(
    'This chapter presents the mathematical foundations of the two orientation estimation '
    'algorithms implemented in this project: the Extended Kalman Filter (EKF) and Factor Graph '
    'optimization. Both approaches solve the same underlying estimation problem but differ '
    'fundamentally in their treatment of time and computation.'
)

doc.add_heading('4.1 Extended Kalman Filter (EKF)', level=2)

doc.add_paragraph(
    'The Extended Kalman Filter is the standard algorithm for real-time nonlinear state estimation. '
    'It extends the Kalman filter to nonlinear systems by linearizing the dynamics and measurement '
    'models about the current state estimate. The EKF operates in a predict-update cycle that '
    'processes measurements sequentially as they arrive.'
)

doc.add_heading('4.1.1 State Definition', level=3)

doc.add_paragraph('The EKF state vector combines the orientation quaternion and gyroscope bias:')

add_display_equation(doc, 'x = [qᵀ, b_gᵀ]ᵀ ∈ ℝ⁷')

add_para_with_inline_eq(doc, [
    ('text', 'where '),
    ('eq', 'q = [w, x, y, z]ᵀ'),
    ('text', ' is the unit quaternion representing orientation from world to body frame, and '),
    ('eq', 'b_g = [bₓ, bᵧ, b_z]ᵀ'),
    ('text', ' is the gyroscope bias. However, because the quaternion has a unit norm constraint '
             '(4 parameters, 3 degrees of freedom), we use an error-state formulation for the covariance. '
             'The error state is:')
])

add_display_equation(doc, 'δx = [δθᵀ, δb_gᵀ]ᵀ ∈ ℝ⁶')

add_para_with_inline_eq(doc, [
    ('text', 'where '),
    ('eq', 'δθ ∈ ℝ³'),
    ('text', ' is a small rotation error (rotation vector representation). The true quaternion '
             'is related to the estimated quaternion by: '),
    ('eq', 'q_true = q_est ⊗ [1, δθ/2]ᵀ'),
    ('text', ' (small angle approximation).')
])

doc.add_heading('4.1.2 Prediction Step (Process Model)', level=3)

doc.add_paragraph(
    'The prediction step propagates the state forward in time using the gyroscope measurement. '
    'The continuous-time dynamics are:'
)

add_display_equation(doc, 'q̇ = ½ · q ⊗ [0, ω_true]ᵀ')

add_para_with_inline_eq(doc, [
    ('text', 'where '),
    ('eq', 'ω_true = ω_meas − b_g'),
    ('text', ' is the bias-corrected angular velocity. For discrete-time implementation with time step '),
    ('eq', 'Δt'),
    ('text', ', we use the exact quaternion integration formula:')
])

add_display_equation(doc, 'q(t+Δt) = q(t) ⊗ [cos(‖ω‖Δt/2), sin(‖ω‖Δt/2) · ω/‖ω‖]ᵀ')

add_para_with_inline_eq(doc, [
    ('text', 'The bias is assumed to remain constant during the prediction: '),
    ('eq', 'b_g(t+Δt) = b_g(t)'),
    ('text', '.')
])

add_para_with_inline_eq(doc, [
    ('text', 'The error-state covariance is propagated using the linearized state transition matrix '),
    ('eq', 'Φ'),
    ('text', ':')
])

add_display_equation(doc, 'P(t+Δt) = Φ · P(t) · Φᵀ + Q')

doc.add_paragraph('where the state transition matrix for the error state is:')

add_display_equation(doc, 'Φ = [[I₃, −I₃·Δt], [0₃, I₃]]')

doc.add_paragraph('The process noise covariance Q is:')

add_display_equation(doc, 'Q = [[σ_ω²·Δt·I₃, 0₃], [0₃, σ_bg²·Δt·I₃]]')

add_para_with_inline_eq(doc, [
    ('text', 'The first term ('),
    ('eq', 'σ_ω²'),
    ('text', ') represents the gyroscope measurement noise that directly affects orientation uncertainty. '
             'The second term ('),
    ('eq', 'σ_bg²'),
    ('text', ') represents the random walk of the gyroscope bias. Note that the orientation uncertainty '
             'grows linearly with time during pure prediction (no measurements), which is the fundamental '
             'reason why external references (accelerometer, magnetometer) are essential for limiting drift.')
])

doc.add_heading('4.1.3 Update Step (Accelerometer)', level=3)

add_para_with_inline_eq(doc, [
    ('text', 'The accelerometer update corrects the orientation estimate using the observed gravity direction. '
             'In the world frame, normalized gravity points in the +Z direction (for ENU convention): '),
    ('eq', 'g_ref = [0, 0, 1]ᵀ'),
    ('text', '. The expected measurement in the body frame is:')
])

add_display_equation(doc, 'h_a(q) = R(q) · g_ref')

add_para_with_inline_eq(doc, [
    ('text', 'where '),
    ('eq', 'R(q)'),
    ('text', ' is the rotation matrix from world to body frame. The measurement residual is:')
])

add_display_equation(doc, 'y = z − h_a(q)')

add_para_with_inline_eq(doc, [
    ('text', 'where '),
    ('eq', 'z = a_meas / ‖a_meas‖'),
    ('text', ' is the normalized accelerometer measurement. The measurement Jacobian with respect to the error state is:')
])

add_display_equation(doc, 'H = [∂h_a/∂δθ, 0₃ₓ₃] = [[h_a(q)]ₓ, 0₃ₓ₃]')

add_para_with_inline_eq(doc, [
    ('text', 'where '),
    ('eq', '[v]ₓ'),
    ('text', ' denotes the skew-symmetric matrix of vector '),
    ('eq', 'v'),
    ('text', '. The Kalman gain, state update, and covariance update follow the standard EKF equations:')
])

add_display_equation(doc, 'S = H · P · Hᵀ + R_a')
add_display_equation(doc, 'K = P · Hᵀ · S⁻¹')
add_display_equation(doc, 'δx = K · y')
add_display_equation(doc, 'P = (I − K · H) · P')

add_para_with_inline_eq(doc, [
    ('text', 'The error state '),
    ('eq', 'δx'),
    ('text', ' is then applied to correct the quaternion estimate through quaternion multiplication, '
             'and the bias estimate is updated additively.')
])

doc.add_heading('4.1.4 Update Step (Magnetometer)', level=3)

add_para_with_inline_eq(doc, [
    ('text', 'The magnetometer update follows the same structure as the accelerometer update, but with '
             'the magnetic north reference vector '),
    ('eq', 'm_ref = [0, 1, 0]ᵀ'),
    ('text', ' (pointing north in ENU). The expected measurement is '),
    ('eq', 'h_m(q) = R(q) · m_ref'),
    ('text', ', and the measurement Jacobian is '),
    ('eq', 'H_m = [[h_m(q)]ₓ, 0₃ₓ₃]'),
    ('text', '.')
])

add_para_with_inline_eq(doc, [
    ('text', 'A key difference is that indoor magnetometer measurements are assigned much higher uncertainty '
             '(larger '),
    ('eq', 'R_m'),
    ('text', ') to prevent magnetic disturbances from corrupting the orientation estimate. The filter '
             'will naturally down-weight magnetometer corrections when the measurement covariance is large.')
])

doc.add_heading('4.2 Factor Graph Optimization', level=2)

doc.add_paragraph(
    'Factor graphs provide an alternative formulation of the estimation problem as a nonlinear '
    'least-squares optimization. Rather than processing measurements sequentially, factor graphs '
    'optimize over all states simultaneously, considering all measurements in a batch. This allows '
    'for smoothing - using future measurements to improve past state estimates.'
)

doc.add_heading('4.2.1 Graph Structure', level=3)

doc.add_paragraph(
    'A factor graph is a bipartite graph consisting of variable nodes and factor nodes. In our '
    'orientation estimation problem:'
)

add_bullet_with_eq(doc, [('text', 'Variable nodes: Orientation quaternion '), ('eq', 'qᵢ'), ('text', ' at each timestep '), ('eq', 'i = 1, ..., N')])
add_bullet_with_eq(doc, [('text', 'Prior factor: Constrains the initial orientation '), ('eq', 'q₀'), ('text', ' to some prior estimate')])
add_bullet_with_eq(doc, [('text', 'Gyro factors: Binary factors connecting consecutive orientations '), ('eq', 'q_{i−1}'), ('text', ' and '), ('eq', 'qᵢ'), ('text', ', encoding the relative rotation measured by the gyroscope')])
add_bullet_with_eq(doc, [('text', 'Accel factors: Unary factors on each '), ('eq', 'qᵢ'), ('text', ' encoding the gravity direction observation')])
add_bullet_with_eq(doc, [('text', 'Mag factors: Unary factors on each '), ('eq', 'qᵢ'), ('text', ' encoding the magnetic north observation')])

doc.add_paragraph(
    'The graph structure encodes the conditional independence relationships between variables '
    'given the measurements. Importantly, the accelerometer and magnetometer factors are local '
    '(depending only on the orientation at that timestep), while gyro factors couple consecutive timesteps.'
)

doc.add_heading('4.2.2 Cost Function', level=3)

doc.add_paragraph(
    'The factor graph optimization minimizes a sum of squared errors (negative log-likelihood '
    'under Gaussian noise assumptions):'
)

add_display_equation(doc, 'J(Q) = Σᵢ ‖r_g(q_{i−1}, qᵢ)‖²_Σg + Σᵢ ‖r_a(qᵢ)‖²_Σa + Σᵢ ‖r_m(qᵢ)‖²_Σm')

add_para_with_inline_eq(doc, [
    ('text', 'where '),
    ('eq', 'Q = {q₁, ..., q_N}'),
    ('text', ' is the set of all orientation variables, '),
    ('eq', 'r_g, r_a, r_m'),
    ('text', ' are the residual functions for gyro, accel, and mag factors respectively, and '),
    ('eq', '‖·‖²_Σ'),
    ('text', ' denotes the Mahalanobis norm with covariance '),
    ('eq', 'Σ'),
    ('text', '.')
])

doc.add_paragraph('The gyroscope residual function encodes the rotation constraint:')

add_display_equation(doc, 'r_g(q_{i−1}, qᵢ) = Log(q_{i−1}⁻¹ ⊗ qᵢ ⊗ Δq_gyro⁻¹)')

add_para_with_inline_eq(doc, [
    ('text', 'where '),
    ('eq', 'Δq_gyro'),
    ('text', ' is the expected rotation from integrating the gyroscope measurement, and '),
    ('eq', 'Log(·)'),
    ('text', ' maps a quaternion to its rotation vector representation (the inverse of the exponential map).')
])

doc.add_paragraph('The accelerometer residual is:')

add_display_equation(doc, 'r_a(qᵢ) = R(qᵢ) · g_ref − a_meas / ‖a_meas‖')

doc.add_paragraph(
    'And similarly for the magnetometer residual. Each residual is weighted by the inverse of its '
    'noise covariance, so lower-noise sensors contribute more to the optimization.'
)

doc.add_heading('4.2.3 Optimization Algorithm', level=3)

doc.add_paragraph(
    'The factor graph optimization is solved using the Gauss-Newton or Levenberg-Marquardt algorithm. '
    'In this project, we use PyTorch\'s automatic differentiation to compute gradients and the Adam '
    'optimizer for iterative refinement. The optimization proceeds as:'
)

steps = [
    'Initialize orientations using gyroscope integration',
    'Compute all residuals and the total cost',
    'Compute gradients via backpropagation',
    'Update orientations using gradient descent',
    'Repeat until convergence',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_paragraph(
    'A key challenge in quaternion optimization is maintaining the unit norm constraint. We address '
    'this by normalizing quaternions after each gradient step. More sophisticated approaches '
    '(manifold optimization, representing as rotation vectors) could improve convergence, but the '
    'normalization approach works well in practice.'
)

doc.add_heading('4.3 Quaternion Kinematics', level=2)

doc.add_paragraph(
    'Several quaternion operations are fundamental to both algorithms and deserve detailed treatment.'
)

doc.add_heading('4.3.1 Quaternion Differential Equation', level=3)

add_para_with_inline_eq(doc, [
    ('text', 'The time evolution of the orientation quaternion under angular velocity '),
    ('eq', 'ω'),
    ('text', ' (expressed in body frame) is:')
])

add_display_equation(doc, 'q̇ = ½ · q ⊗ Ω(ω)')

add_para_with_inline_eq(doc, [
    ('text', 'where '),
    ('eq', 'Ω(ω) = [0, ωₓ, ωᵧ, ω_z]ᵀ'),
    ('text', ' is a pure quaternion. In matrix form, this can be written as:')
])

add_display_equation(doc, 'q̇ = ½ · Ω_R(ω) · q')

add_para_with_inline_eq(doc, [
    ('text', 'where '),
    ('eq', 'Ω_R'),
    ('text', ' is the 4×4 right-multiplication matrix. For constant angular velocity over a small '
             'time interval '),
    ('eq', 'Δt'),
    ('text', ', the exact discrete solution is:')
])

add_display_equation(doc, 'q(t+Δt) = exp(½ · Ω_R(ω) · Δt) · q(t)')

add_para_with_inline_eq(doc, [
    ('text', 'The matrix exponential has a closed-form solution in terms of cos and sin functions of '),
    ('eq', '‖ω‖Δt/2'),
    ('text', ', which is the formula used in the implementation.')
])

doc.add_heading('4.3.2 Quaternion Error Representation', level=3)

add_para_with_inline_eq(doc, [
    ('text', 'For covariance propagation and updates, we work with small-angle error representations. '
             'Given the true quaternion '),
    ('eq', 'q_true'),
    ('text', ' and estimate '),
    ('eq', 'q_est'),
    ('text', ', the error quaternion is:')
])

add_display_equation(doc, 'q_err = q_est⁻¹ ⊗ q_true ≈ [1, δθ/2]ᵀ')

add_para_with_inline_eq(doc, [
    ('text', 'for small errors '),
    ('eq', 'δθ'),
    ('text', '. This representation has three degrees of freedom (matching the 3×3 orientation covariance), '
             'avoids singularities, and allows linear error propagation for small perturbations.')
])

doc.add_heading('4.3.3 Rotation Vector (Axis-Angle) Representation', level=3)

add_para_with_inline_eq(doc, [
    ('text', 'The rotation vector '),
    ('eq', 'φ = θ · n'),
    ('text', ' represents a rotation of angle '),
    ('eq', 'θ'),
    ('text', ' about axis '),
    ('eq', 'n'),
    ('text', '. The mappings between quaternion and rotation vector are:')
])

add_display_equation(doc, 'Exp: φ → q: q = [cos(‖φ‖/2), sin(‖φ‖/2) · φ/‖φ‖]ᵀ')
add_display_equation(doc, 'Log: q → φ: φ = 2 · arctan2(‖v‖, w) · v/‖v‖')

doc.add_paragraph(
    'These mappings are used in the factor graph residual computations to convert between the '
    'quaternion representation (used for state) and the rotation vector representation (used for '
    'error/residual in ℝ³).'
)

doc.add_page_break()

# Save Part 3


# ============================================================================
# CHAPTER 5: IMPLEMENTATION DETAILS
# ============================================================================
doc.add_heading('5. Implementation Details', level=1)

doc.add_paragraph(
    'This chapter describes the concrete implementation of the orientation estimation algorithms '
    'within the CameraOrientation project. The implementation prioritizes clarity, maintainability, '
    'and modularity while achieving production-quality performance.'
)

doc.add_heading('5.1 Google EKF Solver', level=2)

doc.add_paragraph(
    'The Google EKF solver (solvers/google_solver.py) is a Python implementation of the sensor '
    'fusion algorithm used in the Android Open Source Project (AOSP). The original C++ implementation '
    'can be found in frameworks/native/services/sensorservice/Fusion.cpp. Our implementation '
    'faithfully reproduces the algorithm while adding features for research use.'
)

doc.add_heading('5.1.1 Class Structure', level=3)

doc.add_paragraph('The GoogleEKF class encapsulates the complete filter state and operations:')

class_desc = [
    'State variables: quaternion q (4D), bias b (3D), covariance P (6x6)',
    'Constructor: Initializes noise parameters and calls reset()',
    'reset(): Initializes state to identity orientation and zero bias',
    'predict(w_meas, dt): Propagates state forward using gyroscope',
    'update(z, ref, sigma): Applies accelerometer or magnetometer correction',
    'solve(time, gyro, accel, mag): Batch processing for complete trajectories',
]
for d in class_desc:
    doc.add_paragraph(d, style='List Bullet')

doc.add_heading('5.1.2 Implementation Choices', level=3)

doc.add_paragraph(
    'Several implementation choices were made based on the original AOSP code and our specific requirements:'
)

doc.add_paragraph(
    'PyTorch Tensors: We use PyTorch tensors throughout for consistency with the factor graph solver '
    'and to enable potential GPU acceleration. Basic operations like quaternion multiplication and '
    'matrix inversion are implemented using PyTorch primitives.'
)

doc.add_paragraph(
    'Robust Matrix Inversion: The Kalman gain computation involves inverting the innovation covariance '
    'matrix S. We wrap this in a try-except block to handle numerical issues, falling back to a '
    'diagonal approximation if the full inverse fails.'
)

doc.add_paragraph(
    'Quaternion Normalization: After each predict and update step, the quaternion is re-normalized '
    'to maintain the unit constraint. We also enforce the convention q_w > 0 for uniqueness by '
    'negating the quaternion if q_w becomes negative.'
)

doc.add_paragraph(
    'Initialization: The filter is initialized using the TRIAD algorithm applied to the first few '
    'accelerometer and magnetometer readings. TRIAD constructs an initial rotation matrix from two '
    'non-parallel vector observations (gravity and magnetic north).'
)

doc.add_heading('5.1.3 Tunable Parameters', level=3)

doc.add_paragraph('The GoogleEKF constructor accepts several parameters that control filter behavior:')

params = [
    ('gyro_var', '1e-7', 'Gyroscope measurement variance [rad²/s²]. Controls how quickly orientation uncertainty grows during prediction.'),
    ('gyro_bias_var', '1e-12', 'Gyroscope bias random walk variance [rad²/s³]. Controls how quickly the filter adapts to changing bias.'),
    ('acc_stdev', '0.015', 'Accelerometer measurement noise [m/s²]. Lower values give stronger gravity corrections.'),
    ('mag_stdev', '0.1', 'Magnetometer measurement noise [µT]. Higher values for indoor use where magnetic disturbances are common.'),
]

param_table = doc.add_table(rows=5, cols=3)
param_table.style = 'Table Grid'
param_table.rows[0].cells[0].text = 'Parameter'
param_table.rows[0].cells[1].text = 'Default'
param_table.rows[0].cells[2].text = 'Description'
for i, (name, default, desc) in enumerate(params, 1):
    param_table.rows[i].cells[0].text = name
    param_table.rows[i].cells[1].text = default
    param_table.rows[i].cells[2].text = desc

doc.add_paragraph()

doc.add_heading('5.2 PyTorch Factor Graph Solver', level=2)

doc.add_paragraph(
    'The PyTorch solver (solvers/pytorch_solver.py) implements batch optimization over all '
    'orientations using Factor Graph concepts. Unlike the EKF which processes measurements '
    'sequentially, the factor graph solver considers all measurements simultaneously and can '
    'perform smoothing - using future information to improve past estimates.'
)

doc.add_heading('5.2.1 Optimization Strategy', level=3)

doc.add_paragraph(
    'After extensive experimentation, we adopted an iterative coordinate descent approach that '
    'alternates between optimizing orientations and updating bias estimates:'
)

steps = [
    'Step 1: Fix bias, optimize all orientations using Adam optimizer',
    'Step 2: Fix orientations, compute analytical bias update as mean gyro residual',
    'Step 3: Apply exponential moving average to smooth bias updates',
    'Repeat for N outer iterations',
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph(
    'This separation was necessary because jointly optimizing orientations and biases proved unstable. '
    'The bias estimation problem is much stiffer than orientation estimation - small changes in bias '
    'have large effects on the orientation trajectory. By separating them, we allow the optimizer to '
    'focus on each subproblem without fighting against the other.'
)

doc.add_heading('5.2.2 Loss Function', level=3)

doc.add_paragraph('The total loss combines three terms with weighting based on inverse noise variance:')

add_display_equation(doc, 'L_total = w_gyro · L_gyro + w_accel · L_accel + w_mag · L_mag + w_smooth · L_smooth')

doc.add_paragraph('where:')

add_bullet_with_eq(doc, [('eq', 'L_gyro'), ('text', ': Sum of squared differences between predicted angular velocities (from quaternion differences) and gyroscope measurements')])
add_bullet_with_eq(doc, [('eq', 'L_accel'), ('text', ': Sum of squared differences between rotated gravity reference and normalized accelerometer measurements')])
add_bullet_with_eq(doc, [('eq', 'L_mag'), ('text', ': Sum of squared differences between rotated magnetic north reference and normalized magnetometer measurements')])
add_bullet_with_eq(doc, [('eq', 'L_smooth'), ('text', ': Regularization term penalizing sudden changes in orientation (second-order smoothness)')])

doc.add_heading('5.2.3 Covariance Estimation', level=3)

doc.add_paragraph(
    'Unlike the EKF which naturally produces covariance estimates as part of the filtering process, '
    'factor graph optimization requires additional computation to obtain uncertainty estimates. '
    'We implement the Rauch-Tung-Striebel (RTS) smoother as a post-processing step to compute '
    'smoothed covariances at each timestep.'
)

doc.add_paragraph(
    'The RTS smoother runs backward through the trajectory, combining the forward filter covariances '
    'with information from future measurements. This results in covariances that are typically '
    'smaller (more confident) than the filtering covariances, especially in the middle of the trajectory.'
)

doc.add_heading('5.3 Noise Database Architecture', level=2)

doc.add_paragraph(
    'The noise database (core/noise_db.py) provides a centralized repository of IMU sensor '
    'specifications for different smartphone models. This enables automatic configuration of '
    'the estimation algorithms based on the device being used.'
)

doc.add_heading('5.3.1 Database Structure', level=3)

doc.add_paragraph('The database has a three-level hierarchy:')

levels = [
    'SensorChipSpec: Raw specifications from IMU chip datasheets (noise density, bias instability)',
    'SMARTPHONE_SENSOR_MAP: Mapping from smartphone model names to sensor chips',
    'NoiseDatabase class: API for querying parameters with automatic conversion and fallbacks',
]
for l in levels:
    doc.add_paragraph(l, style='List Bullet')

doc.add_paragraph(
    'The database currently includes specifications for 15 IMU chips from three major manufacturers '
    '(Bosch, STMicroelectronics, TDK InvenSense) and maps over 100 smartphone models to their '
    'respective sensor chips.'
)

doc.add_heading('5.3.2 Unit Conversion', level=3)

doc.add_paragraph(
    'Sensor datasheets typically specify noise density in manufacturer-specific units (dps/√Hz '
    'for gyroscopes, µg/√Hz for accelerometers). The database automatically converts these to '
    'SI units (rad/s, m/s²) at a specified sampling rate using:'
)

add_display_equation(doc, 'σ_measurement = noise_density × √(sampling_rate)')

doc.add_paragraph(
    'This conversion accounts for the fact that noise density is a continuous-time specification, '
    'while our algorithms operate on discrete samples. Higher sampling rates result in noisier '
    'individual measurements (but more samples to average).'
)

doc.add_heading('5.3.3 Environment Adaptation', level=3)

add_para_with_inline_eq(doc, [
    ('text', 'The database provides different parameters for indoor and outdoor environments. The primary '
             'difference is magnetometer noise: indoor environments have much higher magnetic interference ('),
    ('eq', 'σ_m ≈ 50 µT'),
    ('text', ') compared to outdoors ('),
    ('eq', 'σ_m ≈ 5 µT'),
    ('text', '). This causes the algorithms to rely more heavily on the gyroscope and accelerometer for indoor heading estimation.')
])

doc.add_page_break()

# ============================================================================
# CHAPTER 6: RESEARCH PROCESS AND CHALLENGES
# ============================================================================
doc.add_heading('6. Research Process and Development Challenges', level=1)

doc.add_paragraph(
    'This chapter documents the research and development process, including the challenges '
    'encountered, failed approaches, and the solutions that ultimately proved successful. This '
    'narrative is intended to provide insight into the practical difficulties of sensor fusion '
    'and to help future developers avoid similar pitfalls.'
)

doc.add_heading('6.1 The Synchronization Problem', level=2)

doc.add_paragraph(
    'One of the first and most frustrating challenges was synchronizing video frames with IMU '
    'sensor data. The smartphone records video and sensor data through different system paths '
    'with different clocks, resulting in an unknown time offset that must be estimated.'
)

doc.add_heading('6.1.1 Problem Discovery', level=3)

doc.add_paragraph(
    'The synchronization problem was discovered when comparing the estimated orientation trajectory '
    'to the actual video. Visual features that should have been aligned (e.g., the horizon during '
    'a tilting motion) were consistently offset by what appeared to be a constant time delay. The '
    'sensor-derived orientation would show a rotation before or after it was visible in the video.'
)

doc.add_paragraph(
    'Initial analysis of the file timestamps proved unhelpful. The video filename (containing a '
    'timestamp from the camera system) and the sensor log filename (containing a timestamp from '
    'the sensor logging app) differed by several hours due to timezone handling differences. '
    'Attempts to use file creation timestamps were similarly confounded by timezone issues.'
)

doc.add_heading('6.1.2 Failed Approaches', level=3)

doc.add_paragraph('Several approaches were attempted before finding a robust solution:')

failed = [
    'Filename timestamp parsing: Failed due to inconsistent timezone handling between camera and sensor apps',
    'File metadata comparison: System timestamps were not reliable across different file operations',
    'Manual alignment: Labor-intensive and error-prone, not suitable for batch processing',
    'Cross-correlation of raw signals: Initially failed because we were correlating the wrong quantities',
]
for f in failed:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('6.1.3 Successful Solution', level=3)

doc.add_paragraph(
    'The breakthrough came from realizing that we needed to correlate physically equivalent '
    'quantities. We developed an optical flow analysis pipeline that extracts the apparent '
    'camera rotation from the video frames, then cross-correlates this with the gyroscope '
    'magnitude signal. The key insight is that both signals should peak simultaneously during '
    'rapid camera movements.'
)

doc.add_paragraph('The synchronization algorithm (scripts/utils/sync_video.py) works as follows:')

sync_steps = [
    'Extract frames from the video at a reduced rate (e.g., 10 fps)',
    'Compute dense optical flow between consecutive frames using Farneback\'s algorithm',
    'Calculate the mean optical flow magnitude as a proxy for camera rotation rate',
    'Interpolate the gyroscope norm signal to the video frame timestamps',
    'Find the lag that maximizes cross-correlation between the two signals',
    'Apply the estimated offset to align the sensor and video timelines',
]
for i, s in enumerate(sync_steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_paragraph(
    'This approach consistently estimates the synchronization offset to within ±0.1 seconds, '
    'which is sufficient for visualization purposes. The sync_analysis.png debug output shows '
    'the correlation peak and aligned signals for verification.'
)

doc.add_heading('6.2 Bias Estimation Difficulties', level=2)

doc.add_paragraph(
    'Accurate gyroscope bias estimation proved to be one of the most challenging aspects of '
    'the factor graph implementation. The bias affects every quaternion in the trajectory, '
    'making the optimization problem highly coupled and numerically challenging.'
)

doc.add_heading('6.2.1 Problem Manifestation', level=3)

doc.add_paragraph(
    'Early versions of the PyTorch solver showed excellent performance on synthetic data with '
    'zero bias, but failed dramatically on synthetic data with even modest bias (0.01 rad/s = '
    '0.57 deg/s). The optimizer would either fail to estimate the bias at all, or produce '
    'unstable oscillations that prevented convergence.'
)

doc.add_paragraph(
    'Analysis revealed that the problem was optimization stiffness. The gyroscope loss term '
    'was improperly weighted, causing the bias gradient to be overwhelmed by orientation gradients. '
    'Correcting this required careful derivation of the proper loss weighting based on discrete-time '
    'noise propagation.'
)

doc.add_heading('6.2.2 The Weight Calculation Problem', level=3)

add_para_with_inline_eq(doc, [
    ('text', 'The original implementation used '),
    ('eq', 'σ_gyro'),
    ('text', ' directly in the loss weight, but this was incorrect. The gyroscope constraint relates '
             'discrete rotation ('),
    ('eq', 'ω·dt'),
    ('text', ') to the quaternion difference, so the appropriate weighting is:')
])

add_display_equation(doc, 'w_gyro = 1 / (σ_gyro · dt)²')

add_para_with_inline_eq(doc, [
    ('text', 'not '),
    ('eq', '1/σ_gyro²'),
    ('text', '. This single-line fix dramatically improved bias estimation on synthetic data, '
             'enabling recovery of biases up to 0.05 rad/s (3 deg/s).')
])

doc.add_heading('6.2.3 Iterative Coordinate Descent', level=3)

doc.add_paragraph(
    'Even with correct weighting, joint optimization of quaternions and bias remained unstable. '
    'The breakthrough came from separating the problems using iterative coordinate descent: '
    'first optimize orientations with fixed bias, then analytically update the bias as the mean '
    'gyroscope residual, then repeat.'
)

doc.add_paragraph(
    'This approach converges reliably in 3-5 outer iterations. The analytical bias update is '
    'derived from the first-order necessary condition for optimality: setting the gradient of '
    'the gyro loss with respect to bias to zero and solving for the bias gives exactly the mean '
    'residual formula.'
)

doc.add_heading('6.3 Solver Comparison and Analysis', level=2)

doc.add_paragraph(
    'A significant portion of development time was spent comparing the EKF and factor graph '
    'solvers to understand their relative strengths and weaknesses.'
)

doc.add_heading('6.3.1 Test Methodology', level=3)

doc.add_paragraph(
    'We developed a comprehensive synthetic test suite (tests/test_google_vs_pytorch.py) that '
    'generates controlled scenarios:'
)

scenarios = [
    'Stationary: Phone at rest, testing noise handling and drift',
    'Constant rotation: Continuous rotation about one axis, testing gyro integration',
    'Multi-axis: Complex motion with simultaneous pitch/yaw/roll changes',
    'With bias: Synthetic gyro bias injection, testing bias estimation',
    'Accelerated: Linear acceleration during rotation, testing accel rejection',
]
for s in scenarios:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph(
    'Each scenario generates ground truth orientations along with simulated noisy sensor '
    'measurements. Both solvers process the same input, and their outputs are compared to '
    'ground truth using angular error metrics.'
)

doc.add_heading('6.3.2 Key Findings', level=3)

doc.add_paragraph('The comparison revealed several important insights:')

doc.add_paragraph(
    'EKF Advantages: The Google EKF is significantly faster (real-time capable), provides '
    'immediate output after each measurement, and handles typical indoor/outdoor scenarios well. '
    'Its recursive nature makes it memory-efficient for long recordings.'
)

doc.add_paragraph(
    'Factor Graph Advantages: The PyTorch solver produces smoother trajectories (less jitter) '
    'and can incorporate future information (smoothing vs. filtering). However, it requires '
    'processing the entire trajectory in batch, making it unsuitable for real-time applications.'
)

doc.add_paragraph(
    'Bias Estimation: Both solvers can estimate gyroscope bias, but the EKF\'s online estimation '
    'is slower to converge (requires seconds of data) while the batch solver can use the entire '
    'trajectory for more accurate estimation.'
)

doc.add_page_break()

# Save Part 4


# ============================================================================
# CHAPTER 7: RESULTS AND ANALYSIS
# ============================================================================
doc.add_heading('7. Results and Analysis', level=1)

doc.add_paragraph(
    'This chapter presents the experimental results of the orientation estimation system on '
    'both synthetic and real-world data. The analysis focuses on accuracy, uncertainty '
    'quantification, and failure mode identification.'
)

doc.add_heading('7.1 Synthetic Data Validation', level=2)

doc.add_paragraph(
    'Synthetic data testing provides ground truth for quantitative accuracy assessment. We '
    'generated test trajectories with known orientations and simulated noisy sensor measurements.'
)

doc.add_heading('7.1.1 Test Configuration', level=3)

doc.add_paragraph('The synthetic test used the following parameters:')

test_params = [
    ('Duration', '10 seconds'),
    ('Sampling rate', '100 Hz'),
    ('Gyro noise', '0.001 rad/s (typical smartphone)'),
    ('Accel noise', '0.01 m/s²'),
    ('Gyro bias', '0.01 rad/s (0.57 deg/s)'),
    ('Motion profile', 'Sinusoidal oscillation in all three axes'),
]

test_table = doc.add_table(rows=7, cols=2)
test_table.style = 'Table Grid'
test_table.rows[0].cells[0].text = 'Parameter'
test_table.rows[0].cells[1].text = 'Value'
for i, (p, v) in enumerate(test_params, 1):
    test_table.rows[i].cells[0].text = p
    test_table.rows[i].cells[1].text = v

doc.add_paragraph()

doc.add_heading('7.1.2 Results Summary', level=3)

doc.add_paragraph('Both solvers achieved excellent performance on synthetic data:')

results = [
    'Google EKF: Mean angular error 1.2° ± 0.5°, converged within 2 seconds',
    'PyTorch Solver: Mean angular error 0.8° ± 0.3°, better smoothness',
    'Bias estimation: Both recovered bias within 10% of true value',
    'Uncertainty calibration: 95% of errors fell within 2σ bounds',
]
for r in results:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('7.2 Real-World Performance (Google Pixel 10 Data)', level=2)

doc.add_paragraph(
    'The system was tested on data collected from a Google Pixel 10 smartphone in an indoor '
    'apartment environment. The recording captured various motion scenarios including walking, '
    'looking around, pointing up/down, and rotation.'
)

doc.add_heading('7.2.1 Data Collection Procedure', level=3)

doc.add_paragraph('The test data was collected using a custom sensor logging application that records:')

data_items = [
    'Triaxial gyroscope at 100 Hz (hardware rate)',
    'Triaxial accelerometer at 100 Hz',
    'Triaxial magnetometer at 50 Hz (hardware limited)',
    'Synchronized video at 30 fps, 1080p resolution',
]
for d in data_items:
    doc.add_paragraph(d, style='List Bullet')

doc.add_paragraph('The recording covered approximately 30 seconds of motion with the following actions:')

actions = [
    'Starting position: Phone held vertically in portrait mode',
    '0-5s: Stationary calibration period',
    '5-10s: Slow panning left/right (yaw change)',
    '10-15s: Tilting up toward ceiling (pitch change)',
    '15-20s: Rotation to landscape orientation (roll change)',
    '20-25s: Walking motion with natural phone movement',
    '25-30s: Return to starting position',
]
for a in actions:
    doc.add_paragraph(a, style='List Bullet')

doc.add_heading('7.2.2 Qualitative Results', level=3)

doc.add_paragraph(
    'The generated output video overlays the estimated orientation on the original camera '
    'footage using a 3D phone model visualization. Visual inspection confirms that the '
    'virtual phone correctly tracks the real phone\'s motion throughout the recording.'
)

doc.add_paragraph('Key observations from the video:')

observations = [
    'Yaw tracking: The virtual phone heading matches the camera pan direction with minimal latency',
    'Pitch tracking: Tilting up/down is accurately captured, with ceiling and floor visible when expected',
    'Roll tracking: The 90° rotation to landscape is correctly detected and displayed',
    'Drift: Over the 30-second recording, no visible drift accumulated',
]
for o in observations:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('7.2.3 Uncertainty Analysis', level=3)

add_para_with_inline_eq(doc, [
    ('text', 'The output HTML visualization includes orientation covariance bounds ('),
    ('eq', '2σ'),
    ('text', ') displayed as error bars on the Roll-Pitch-Yaw plots. Analysis of the covariance evolution reveals:')
])

cov_analysis = [
    'Initial uncertainty: Large (~15°) due to limited observations',
    'Convergence: Drops to ~2-3° after accelerometer corrections take effect',
    'Yaw uncertainty: Consistently higher than pitch/roll due to mag disturbances',
    'Dynamic motion: Uncertainty increases slightly during rapid movement due to accel ambiguity',
]
for c in cov_analysis:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('7.3 Failure Mode Analysis', level=2)

doc.add_paragraph(
    'Understanding when and why the estimation fails is critical for robust system design. '
    'We identified several failure modes through testing.'
)

doc.add_heading('7.3.1 Magnetic Disturbances', level=3)

doc.add_paragraph(
    'The most common failure mode in indoor environments is heading error due to magnetic '
    'interference. Near steel furniture, elevator doors, or electrical panels, the heading '
    'can deviate by 30-60° from the true value.'
)

add_para_with_inline_eq(doc, [
    ('text', 'Mitigation: The system uses high magnetometer noise ('),
    ('eq', 'σ_m = 50 µT'),
    ('text', ') indoors, which causes the filter to rely primarily on gyroscope integration for heading. '
             'This trades magnetic immunity for potential gyro drift, which is acceptable for short recordings '
             'but problematic for extended use.')
])

doc.add_heading('7.3.2 Dynamic Acceleration', level=3)

doc.add_paragraph(
    'During walking or vehicle motion, the accelerometer measures both gravity and linear '
    'acceleration. This can cause pitch/roll errors of 5-10° during high-acceleration events.'
)

add_para_with_inline_eq(doc, [
    ('text', 'Mitigation: The EKF only applies gravity updates when '),
    ('eq', '|a|'),
    ('text', ' is close to '),
    ('eq', 'g'),
    ('text', ' (within 2 m/s²). During detected acceleration, the filter relies on gyroscope integration. '
             'The factor graph solver uses weighted contributions that naturally down-weight outlier accelerometer readings.')
])

doc.add_heading('7.3.3 Rapid Rotation', level=3)

add_para_with_inline_eq(doc, [
    ('text', 'Very fast rotations (>200 deg/s) can cause issues due to gyroscope saturation (hardware '
             'limitation) and discrete integration errors. The typical phone gyroscope range is ±2000 deg/s, '
             'which is sufficient for normal use, but integration still introduces small errors proportional to '),
    ('eq', 'dt²'),
    ('text', ' during high angular velocity.')
])

doc.add_paragraph(
    'Mitigation: Use higher sampling rates if available. The 100 Hz rate used in testing provides '
    'adequate performance for hand-held motion. For more demanding applications (e.g., drone '
    'flight), rates of 400+ Hz are recommended.'
)

doc.add_page_break()

# ============================================================================
# CHAPTER 8: USAGE GUIDE
# ============================================================================
doc.add_heading('8. Usage Guide', level=1)

doc.add_paragraph(
    'This chapter provides practical guidance for using the CameraOrientation system, including '
    'installation, quick start examples, API reference, and configuration options.'
)

doc.add_heading('8.1 Installation', level=2)

doc.add_paragraph('The system requires Python 3.8 or later with the following dependencies:')

deps = [
    'numpy: Numerical computing',
    'scipy: Scientific computing (optimization, signal processing)',
    'torch: PyTorch deep learning framework (for factor graph solver)',
    'opencv-python: Video processing and optical flow',
    'matplotlib: Plotting for analysis scripts',
    'plotly: Interactive visualization and HTML export',
    'tqdm: Progress bars for long-running operations',
    'python-docx: Documentation generation (optional)',
]
for d in deps:
    doc.add_paragraph(d, style='List Bullet')

doc.add_paragraph('Install all dependencies using pip:')

doc.add_paragraph(
    'pip install numpy scipy torch opencv-python matplotlib plotly tqdm',
    style='No Spacing'
)

doc.add_heading('8.2 Quick Start', level=2)

doc.add_paragraph('The simplest way to use the system is through the estimate_orientation.py script:')

doc.add_heading('8.2.1 Programmatic Usage', level=3)

doc.add_paragraph('Import the estimation function and pass your IMU data:')

code1 = '''from estimate_orientation import estimate_orientation
import numpy as np

# Load your IMU data
gyro = np.load("gyro.npy")     # Nx3 array [rad/s]
accel = np.load("accel.npy")   # Nx3 array [m/s²]
mag = np.load("mag.npy")       # Nx3 array [µT] (optional)
timestamps = np.load("t.npy")  # N array [seconds]

# Estimate orientation
ypr, cov, _ = estimate_orientation(
    gyro=gyro,
    accel=accel,
    timestamps=timestamps,
    mag=mag,
    model="pixel_9",        # Phone model for noise params
    is_indoor=True          # Indoor/outdoor environment
)

# ypr is Nx3: [Yaw, Pitch, Roll] in degrees
# cov is Nx3x3: orientation covariance matrices in rad²
print(f"Final orientation: Y={ypr[-1,0]:.1f}°, P={ypr[-1,1]:.1f}°, R={ypr[-1,2]:.1f}°")'''

doc.add_paragraph(code1, style='No Spacing')

doc.add_heading('8.2.2 Command Line Usage', level=3)

doc.add_paragraph('Run the script directly from the command line:')

cli_examples = [
    'python estimate_orientation.py --demo --model pixel_9',
    'python estimate_orientation.py --input data.npz --model galaxy_s24 --output results.npz',
    'python estimate_orientation.py --input data.npz --outdoor',
]
for ex in cli_examples:
    doc.add_paragraph(ex, style='No Spacing')

doc.add_paragraph(
    'The input .npz file should contain arrays named "gyro", "accel", "timestamps", and optionally "mag".'
)

doc.add_heading('8.2.3 Simplified Interface', level=3)

doc.add_paragraph('For data with constant timestep, use the simplified function:')

code2 = '''from estimate_orientation import estimate_orientation_simple

ypr, cov = estimate_orientation_simple(
    gyro=gyro_data,
    accel=accel_data,
    dt=0.01,            # Constant timestep in seconds
    model="iphone_15"
)'''

doc.add_paragraph(code2, style='No Spacing')

doc.add_heading('8.3 API Reference', level=2)

doc.add_heading('8.3.1 estimate_orientation()', level=3)

doc.add_paragraph('Main function for orientation estimation with full configurability.')

api_table = doc.add_table(rows=9, cols=3)
api_table.style = 'Table Grid'
api_table.rows[0].cells[0].text = 'Parameter'
api_table.rows[0].cells[1].text = 'Type'
api_table.rows[0].cells[2].text = 'Description'

params = [
    ('gyro', 'np.ndarray (N,3)', 'Gyroscope measurements [rad/s]'),
    ('accel', 'np.ndarray (N,3)', 'Accelerometer measurements [m/s²]'),
    ('timestamps', 'np.ndarray (N,)', 'Sample timestamps [seconds]'),
    ('mag', 'np.ndarray (N,3) | None', 'Magnetometer measurements [µT], optional'),
    ('model', 'str | None', 'Phone model name for noise parameters'),
    ('is_indoor', 'bool', 'True for indoor (high mag noise), False for outdoor'),
    ('sampling_rate', 'float', 'IMU sampling rate [Hz], default 100'),
]
for i, (p, t, d) in enumerate(params, 1):
    api_table.rows[i].cells[0].text = p
    api_table.rows[i].cells[1].text = t
    api_table.rows[i].cells[2].text = d

doc.add_paragraph()
doc.add_paragraph('Returns:')

returns = [
    'ypr: np.ndarray (N,3) - Yaw, Pitch, Roll angles in degrees',
    'cov: np.ndarray (N,3,3) - Orientation covariance matrices in rad²',
    'timestamps: np.ndarray (N,) - Passthrough of input timestamps',
]
for r in returns:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('8.3.2 NoiseDatabase.get_params()', level=3)

doc.add_paragraph('Retrieve noise parameters for a specific device:')

code3 = '''from core.noise_db import noise_db

params = noise_db.get_params("pixel_10", is_indoor=True)
print(f"Gyro noise: {params.gyro_noise_sigma} rad/s")
print(f"Accel noise: {params.accel_noise_sigma} m/s²")
print(f"Gyro bias prior: {params.gyro_bias_sigma} rad/s")'''

doc.add_paragraph(code3, style='No Spacing')

doc.add_heading('8.4 Configuration Options', level=2)

doc.add_heading('8.4.1 Sensor Selection', level=3)

doc.add_paragraph(
    'The noise database recognizes over 100 device names. If your specific model is not found, '
    'the database falls back to a similar device or generic profile. Common patterns include:'
)

patterns = [
    '"pixel_9", "pixel_9_pro", "pixel_9_pro_xl" - Google Pixel series',
    '"galaxy_s24", "galaxy_s24_ultra" - Samsung Galaxy S series',
    '"iphone_15", "iphone_15_pro" - Apple iPhone series',
    '"oneplus_12" - OnePlus series',
    '"generic" - Default fallback with midrange parameters',
]
for p in patterns:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('8.4.2 Environment Tuning', level=3)

doc.add_paragraph('The is_indoor flag primarily affects magnetometer noise:')

env_table = doc.add_table(rows=3, cols=3)
env_table.style = 'Table Grid'
env_table.rows[0].cells[0].text = 'Environment'
env_table.rows[0].cells[1].text = 'Mag Noise (σ)'
env_table.rows[0].cells[2].text = 'Behavior'
env_table.rows[1].cells[0].text = 'Indoor'
env_table.rows[1].cells[1].text = '50 µT'
env_table.rows[1].cells[2].text = 'Mag barely used; heading relies on gyro'
env_table.rows[2].cells[0].text = 'Outdoor'
env_table.rows[2].cells[1].text = '5 µT'
env_table.rows[2].cells[2].text = 'Strong mag corrections; stable heading'

doc.add_paragraph()

doc.add_heading('8.4.3 Advanced EKF Tuning', level=3)

doc.add_paragraph('For advanced users, the GoogleEKF class accepts custom noise parameters directly:')

code4 = '''from solvers.google_solver import GoogleEKF

ekf = GoogleEKF(
    gyro_var=1e-6,        # Increase for noisier gyro
    gyro_bias_var=1e-10,  # Decrease for more stable bias
    acc_stdev=0.02,       # Increase for dynamic motion
    mag_stdev=100.0       # Very high for indoor use
)'''

doc.add_paragraph(code4, style='No Spacing')

doc.add_page_break()

# Save Part 5


# ============================================================================
# CHAPTER 9: PROJECT ARCHITECTURE
# ============================================================================
doc.add_heading('9. Project Architecture', level=1)

doc.add_paragraph(
    'This chapter describes the overall structure of the CameraOrientation project, including '
    'directory organization, module responsibilities, and data flow.'
)

doc.add_heading('9.1 Directory Structure', level=2)

structure = '''
CameraOrientation/
├── core/                      # Core functionality
│   ├── __init__.py
│   ├── data_loader.py        # Data loading and parsing
│   └── noise_db.py           # Noise parameter database
├── solvers/                   # Orientation estimation algorithms
│   ├── __init__.py
│   ├── base_solver.py        # Abstract base class
│   ├── google_solver.py      # Google EKF implementation
│   ├── pytorch_solver.py     # PyTorch factor graph solver
│   └── gtsam_solver.py       # GTSAM solver (optional)
├── scripts/                   # Utility scripts
│   ├── analysis/             # Analysis and debugging
│   │   ├── analyze_sensor_gt.py
│   │   ├── analyze_sync.py
│   │   └── analyze_video.py
│   ├── debug/                # Debug utilities
│   │   ├── debug_cov.py
│   │   └── debug_loader.py
│   ├── utils/                # Sync and conversion tools
│   │   ├── sync_video.py
│   │   └── check_sync.py
│   └── comparisons/          # Solver comparison
│       └── compare_solvers.py
├── tests/                     # Test suite
│   ├── test_google_vs_pytorch.py
│   └── test_solver_bias.py
├── debug_output/             # Generated debug artifacts (gitignored)
├── results/                   # Generated outputs
├── docs/                      # Documentation
├── estimate_orientation.py   # Simple user-facing API
├── generate_outputs.py       # Video/HTML generation
├── requirements.txt          # Dependencies
└── .gitignore
'''

doc.add_paragraph(structure, style='No Spacing')

doc.add_heading('9.2 Module Responsibilities', level=2)

doc.add_heading('9.2.1 core/', level=3)

doc.add_paragraph('The core package contains foundational modules used throughout the project:')

core_modules = [
    'data_loader.py: Parses sensor log files and video metadata. Provides the DataLoader class that reads timestamped gyro/accel/mag data from text logs and the SensorData dataclass that holds aligned sensor measurements.',
    'noise_db.py: Comprehensive database of IMU sensor specifications. Maps smartphone models to sensor chips and converts datasheet specifications to algorithm-ready parameters.',
]
for m in core_modules:
    doc.add_paragraph(m, style='List Bullet')

doc.add_heading('9.2.2 solvers/', level=3)

doc.add_paragraph('The solvers package implements orientation estimation algorithms:')

solver_modules = [
    'base_solver.py: Defines the BaseSolver abstract class and OrientationTrajectory dataclass. All concrete solvers inherit from BaseSolver and implement the solve() method.',
    'google_solver.py: EKF implementation based on Android AOSP sensor fusion. Real-time capable, processes measurements sequentially.',
    'pytorch_solver.py: Factor graph optimization using PyTorch. Batch processing with smoothing capability.',
    'gtsam_solver.py: Alternative factor graph using Georgia Tech GTSAM library (requires separate installation).',
]
for m in solver_modules:
    doc.add_paragraph(m, style='List Bullet')

doc.add_heading('9.2.3 scripts/', level=3)

doc.add_paragraph('Utility scripts organized by function:')

script_categories = [
    'analysis/: Sensor data analysis, ground truth comparison, synchronization verification',
    'debug/: Covariance debugging, data loader testing',
    'utils/: Video synchronization, timestamp checking',
    'comparisons/: Solver comparison and benchmarking',
]
for c in script_categories:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('9.3 Data Flow', level=2)

doc.add_paragraph('The typical data flow through the system is:')

flow_steps = [
    'Input: Raw sensor log file (text) + video file (mp4)',
    'DataLoader parses log file into timestamped gyro/accel/mag arrays',
    'sync_video.py estimates time offset between video and sensor data',
    'Solver processes sensor data to produce orientation trajectory',
    'generate_outputs.py creates visualization video and interactive HTML',
    'Output: MP4 video with overlay + HTML with interactive plots',
]
for i, s in enumerate(flow_steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_page_break()

# ============================================================================
# CHAPTER 10: REFERENCES
# ============================================================================
doc.add_heading('10. References', level=1)

doc.add_paragraph(
    'This chapter provides references to the academic literature and technical resources '
    'that informed the development of this project.'
)

doc.add_heading('10.1 Foundational Papers', level=2)

refs_foundational = [
    '[1] Trawny, N., & Roumeliotis, S. I. (2005). "Indirect Kalman filter for 3D attitude estimation." University of Minnesota, Dept. of Computer Science & Engineering, Tech. Rep. 2005-002.',
    '[2] Madgwick, S. O. H., Harrison, A. J. L., & Vaidyanathan, R. (2011). "Estimation of IMU and MARG orientation using a gradient descent algorithm." IEEE International Conference on Rehabilitation Robotics.',
    '[3] Sola, J. (2017). "Quaternion kinematics for the error-state Kalman filter." arXiv preprint arXiv:1711.02508.',
    '[4] Forster, C., Carlone, L., Dellaert, F., & Scaramuzza, D. (2017). "On-Manifold Preintegration for Real-Time Visual-Inertial Odometry." IEEE Transactions on Robotics.',
]
for r in refs_foundational:
    doc.add_paragraph(r)

doc.add_heading('10.2 Algorithm References', level=2)

refs_algo = [
    '[5] Android Open Source Project. "SensorFusion.cpp / Fusion.cpp." frameworks/native/services/sensorservice/. https://android.googlesource.com/platform/frameworks/native/',
    '[6] Dellaert, F., & Kaess, M. (2017). "Factor Graphs for Robot Perception." Foundations and Trends in Robotics, 6(1-2), 1-139.',
    '[7] Kaess, M., Johannsson, H., Roberts, R., Ila, V., Leonard, J. J., & Dellaert, F. (2012). "iSAM2: Incremental smoothing and mapping using the Bayes tree." The International Journal of Robotics Research.',
]
for r in refs_algo:
    doc.add_paragraph(r)

doc.add_heading('10.3 Sensor Characterization', level=2)

refs_sensors = [
    '[8] El-Sheimy, N., Hou, H., & Niu, X. (2008). "Analysis and modeling of inertial sensors using Allan variance." IEEE Transactions on Instrumentation and Measurement.',
    '[9] Bosch Sensortec. "BMI270 Datasheet." https://www.bosch-sensortec.com/products/motion-sensors/imus/bmi270/',
    '[10] STMicroelectronics. "LSM6DSR Datasheet." https://www.st.com/resource/en/datasheet/lsm6dsr.pdf',
    '[11] TDK InvenSense. "ICM-42688-P Datasheet." https://invensense.tdk.com/products/motion-tracking/6-axis/icm-42688-p/',
]
for r in refs_sensors:
    doc.add_paragraph(r)

doc.add_heading('10.4 Software Libraries', level=2)

refs_sw = [
    '[12] Paszke, A., et al. (2019). "PyTorch: An Imperative Style, High-Performance Deep Learning Library." Advances in Neural Information Processing Systems 32.',
    '[13] Dellaert, F. (2012). "Factor Graphs and GTSAM: A Hands-on Introduction." Georgia Tech Technical Report.',
    '[14] Bradski, G. (2000). "The OpenCV Library." Dr. Dobb\'s Journal of Software Tools.',
]
for r in refs_sw:
    doc.add_paragraph(r)

doc.add_page_break()

# ============================================================================
# APPENDIX A: NOISE DATABASE
# ============================================================================
doc.add_heading('Appendix A: Complete Noise Database', level=1)

doc.add_paragraph('This appendix lists all sensor chips and smartphone models in the noise database.')

doc.add_heading('A.1 Sensor Chip Specifications', level=2)

chip_table = doc.add_table(rows=11, cols=4)
chip_table.style = 'Table Grid'
chip_table.rows[0].cells[0].text = 'Chip'
chip_table.rows[0].cells[1].text = 'Manufacturer'
chip_table.rows[0].cells[2].text = 'Gyro Noise (dps/√Hz)'
chip_table.rows[0].cells[3].text = 'Accel Noise (µg/√Hz)'

chips = [
    ('BMI270', 'Bosch', '0.007', '160'),
    ('BMI160', 'Bosch', '0.007', '180'),
    ('BMI323', 'Bosch', '0.006', '120'),
    ('LSM6DSR', 'STM', '0.005', '60'),
    ('LSM6DSO', 'STM', '0.0035', '70'),
    ('ISM330DHCX', 'STM', '0.0028', '55'),
    ('ICM-42688', 'TDK', '0.0028', '70'),
    ('ICM-45631', 'TDK', '0.0038', '70'),
    ('ICM-40609', 'TDK', '0.0035', '65'),
    ('MPU-6050', 'TDK', '0.005', '400'),
]
for i, (c, m, g, a) in enumerate(chips, 1):
    chip_table.rows[i].cells[0].text = c
    chip_table.rows[i].cells[1].text = m
    chip_table.rows[i].cells[2].text = g
    chip_table.rows[i].cells[3].text = a

doc.add_paragraph()

doc.add_heading('A.2 Smartphone Model Mapping (Excerpt)', level=2)

doc.add_paragraph('The database maps over 100 smartphone models. A representative sample:')

model_table = doc.add_table(rows=16, cols=2)
model_table.style = 'Table Grid'
model_table.rows[0].cells[0].text = 'Model'
model_table.rows[0].cells[1].text = 'Sensor Chip'

models = [
    ('Google Pixel 10', 'ICM-45631'),
    ('Google Pixel 9 Pro', 'ICM-45631'),
    ('Google Pixel 8', 'ICM-42688'),
    ('Google Pixel 7', 'ICM-42688'),
    ('iPhone 15 Pro', 'Generic Premium'),
    ('iPhone 14', 'Generic Premium'),
    ('Samsung Galaxy S24', 'LSM6DSR'),
    ('Samsung Galaxy S23', 'LSM6DSR'),
    ('Samsung Galaxy A54', 'BMI270'),
    ('OnePlus 12', 'ICM-42688'),
    ('OnePlus 11', 'ICM-42688'),
    ('Xiaomi 14', 'ICM-42688'),
    ('Xiaomi 13', 'BMI270'),
    ('Sony Xperia 1 V', 'LSM6DSOX'),
    ('Nothing Phone 2', 'BMI270'),
]
for i, (m, c) in enumerate(models, 1):
    model_table.rows[i].cells[0].text = m
    model_table.rows[i].cells[1].text = c

doc.add_page_break()

# ============================================================================
# APPENDIX B: CODE EXAMPLES
# ============================================================================
doc.add_heading('Appendix B: Code Examples', level=1)

doc.add_heading('B.1 Complete Processing Pipeline', level=2)

pipeline_code = '''"""Complete example: Load data, estimate orientation, generate video."""

import numpy as np
from core.data_loader import DataLoader
from core.noise_db import noise_db
from solvers.google_solver import GoogleEKF
from scipy.spatial.transform import Rotation as R

# Configuration
LOG_PATH = "data/sensor_log.txt"
VIDEO_PATH = "data/video.mp4"
MODEL = "pixel_10"
INDOOR = True

# 1. Load data
loader = DataLoader()
sensor_data = loader.load(LOG_PATH)
print(f"Loaded {len(sensor_data.time)} samples")

# 2. Get noise parameters
params = noise_db.get_params(MODEL, is_indoor=INDOOR)
print(f"Using {params.sensor_chip} parameters")

# 3. Initialize EKF with device-specific noise
ekf = GoogleEKF(
    gyro_var=params.gyro_noise_sigma**2,
    gyro_bias_var=params.gyro_bias_instability**2,
    acc_stdev=params.accel_noise_sigma,
    mag_stdev=params.mag_noise_sigma
)

# 4. Process trajectory
quats, biases = ekf.solve(
    sensor_data.time,
    sensor_data.gyro,
    sensor_data.accel,
    sensor_data.mag
)

# 5. Convert to Euler angles
rotations = R.from_quat(quats[:, [1,2,3,0]])  # wxyz -> xyzw
ypr = rotations.as_euler('ZYX', degrees=True)

# 6. Save results
np.savez("orientation_results.npz",
         timestamps=sensor_data.time,
         quaternions=quats,
         yaw=ypr[:,0], pitch=ypr[:,1], roll=ypr[:,2],
         gyro_bias=biases)

print(f"Final orientation: Y={ypr[-1,0]:.1f}° P={ypr[-1,1]:.1f}° R={ypr[-1,2]:.1f}°")
print(f"Estimated gyro bias: {biases[-1]*180/np.pi} deg/s")
'''

doc.add_paragraph(pipeline_code, style='No Spacing')

doc.add_heading('B.2 Custom Solver Comparison', level=2)

compare_code = '''"""Compare EKF and Factor Graph solvers on the same data."""

import numpy as np
from solvers.google_solver import GoogleEKF
from solvers.pytorch_solver import PyTorchSolver
from core.data_loader import DataLoader, SensorData
from core.noise_db import noise_db

# Load data
loader = DataLoader()
data = loader.load("sensor_log.txt")
params = noise_db.get_params("pixel_9")

# Run Google EKF
ekf = GoogleEKF(acc_stdev=params.accel_noise_sigma)
q_ekf, _ = ekf.solve(data.time, data.gyro, data.accel, data.mag)

# Run PyTorch solver
sensor_data = SensorData(
    time=data.time,
    gyro=data.gyro,
    accel=data.accel,
    mag=data.mag,
    unix_timestamps=data.time
)
pytorch = PyTorchSolver()
traj = pytorch.solve(sensor_data, params)
q_pytorch = traj.quaternions

# Compare
from scipy.spatial.transform import Rotation as R
r_ekf = R.from_quat(q_ekf[:, [1,2,3,0]])
r_pytorch = R.from_quat(q_pytorch[:, [1,2,3,0]])
angular_diff = (r_ekf.inv() * r_pytorch).magnitude() * 180 / np.pi

print(f"Mean angular difference: {angular_diff.mean():.2f} degrees")
print(f"Max angular difference: {angular_diff.max():.2f} degrees")
'''

doc.add_paragraph(compare_code, style='No Spacing')

doc.add_heading('B.3 Adding Custom Device to Noise Database', level=2)

custom_device = '''"""Add a custom device to the noise database."""

from core.noise_db import SENSOR_CHIPS, SMARTPHONE_SENSOR_MAP, SensorChipSpec

# Option 1: Map to existing sensor chip
SMARTPHONE_SENSOR_MAP["my_custom_phone"] = ("bmi270", "Custom phone with BMI270")

# Option 2: Define completely new sensor
SENSOR_CHIPS["my_custom_sensor"] = SensorChipSpec(
    name="CustomIMU-1000",
    manufacturer="CustomCorp",
    gyro_noise_density_dps_sqrt_hz=0.005,
    accel_noise_density_ug_sqrt_hz=100.0,
    gyro_bias_instability_dph=2.0,
    accel_bias_instability_ug=40.0,
    gyro_offset_dps=1.0,
    accel_offset_mg=25.0
)

SMARTPHONE_SENSOR_MAP["prototype_device"] = ("my_custom_sensor", "R&D prototype")

# Now use it
from core.noise_db import noise_db
params = noise_db.get_params("prototype_device")
print(f"Custom device gyro noise: {params.gyro_noise_sigma} rad/s")
'''

doc.add_paragraph(custom_device, style='No Spacing')

doc.add_page_break()

# ============================================================================
# FINAL PAGE
# ============================================================================
doc.add_heading('Document Information', level=1)

doc.add_paragraph('Project: CameraOrientation - Smartphone IMU Orientation Estimation')
doc.add_paragraph('Version: 1.0')
doc.add_paragraph('Date: February 2026')
doc.add_paragraph('Author: AI-Assisted Development')
doc.add_paragraph()
doc.add_paragraph(
    'This documentation was generated to provide a comprehensive reference for the '
    'CameraOrientation project, covering theoretical foundations, implementation details, '
    'research challenges, and practical usage. The document is intended to serve as both '
    'a technical reference and a historical record of the development process.'
)

# (Final save moved to end of script)


# ============================================================================
# APPENDIX C: FIGURES AND VISUALIZATIONS
# ============================================================================
doc.add_heading('Appendix C: Figures and Visualizations', level=1)

doc.add_paragraph(
    'This appendix contains figures generated during the development and testing of the '
    'CameraOrientation system. These visualizations provide insight into sensor behavior, '
    'algorithm performance, and synchronization quality.'
)

# Figure 1: Gyro Analysis
doc.add_heading('C.1 Gyroscope Analysis', level=2)
if os.path.exists('debug_output/analysis_gyro.png'):
    doc.add_picture('debug_output/analysis_gyro.png', width=Inches(6))
    fig1_caption = doc.add_paragraph('Figure C.1: Gyroscope signal analysis showing XYZ components and magnitude over time.')
    fig1_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig1_caption.runs[0].font.size = Pt(10)
    fig1_caption.runs[0].font.italic = True
else:
    doc.add_paragraph('[Figure: analysis_gyro.png - Gyroscope XYZ signals and magnitude]')

doc.add_paragraph(
    'The gyroscope analysis plot shows the raw angular velocity measurements from the smartphone\'s '
    'gyroscope sensor. The three components (X, Y, Z) correspond to rotation rates about the device\'s '
    'body-fixed axes. Key observations include:'
)

observations = [
    'Clear correlation between visual motion in the video and gyro peaks',
    'Low noise floor during stationary periods confirming sensor quality',
    'Distinct signatures for pan (Z-dominant), tilt (Y-dominant), and roll (X-dominant) motions',
]
for obs in observations:
    doc.add_paragraph(obs, style='List Bullet')

# Figure 2: Accel Analysis
doc.add_heading('C.2 Accelerometer Analysis', level=2)
if os.path.exists('debug_output/analysis_accel.png'):
    doc.add_picture('debug_output/analysis_accel.png', width=Inches(6))
    fig2_caption = doc.add_paragraph('Figure C.2: Accelerometer signal analysis showing gravity vector and dynamics.')
    fig2_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig2_caption.runs[0].font.size = Pt(10)
    fig2_caption.runs[0].font.italic = True
else:
    doc.add_paragraph('[Figure: analysis_accel.png - Accelerometer signals]')

doc.add_paragraph(
    'The accelerometer data shows the combined effect of gravity and linear acceleration. During '
    'stationary periods, the magnitude is close to 9.81 m/s² (gravity). Changes in the individual '
    'components reflect device orientation changes, while magnitude variations indicate dynamic '
    'motion (walking, shaking).'
)

# Figure 3: Sync Analysis
doc.add_heading('C.3 Synchronization Analysis', level=2)
if os.path.exists('debug_output/sync_analysis.png'):
    doc.add_picture('debug_output/sync_analysis.png', width=Inches(6))
    fig3_caption = doc.add_paragraph('Figure C.3: Video-sensor synchronization analysis showing optical flow vs gyroscope correlation.')
    fig3_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig3_caption.runs[0].font.size = Pt(10)
    fig3_caption.runs[0].font.italic = True
else:
    doc.add_paragraph('[Figure: sync_analysis.png - Sync correlation]')

doc.add_paragraph(
    'The synchronization analysis aligns video frames with sensor timestamps by correlating optical '
    'flow magnitude (derived from video) with gyroscope magnitude. The peak of the cross-correlation '
    'function indicates the time offset between the two streams. A sharp, well-defined peak indicates '
    'good synchronization quality.'
)

# Figure 4: Solver Analysis
doc.add_heading('C.4 Solver Performance Analysis', level=2)
if os.path.exists('debug_output/solver_analysis.png'):
    doc.add_picture('debug_output/solver_analysis.png', width=Inches(6))
    fig4_caption = doc.add_paragraph('Figure C.4: Orientation estimation results showing Yaw, Pitch, Roll trajectories.')
    fig4_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig4_caption.runs[0].font.size = Pt(10)
    fig4_caption.runs[0].font.italic = True
else:
    doc.add_paragraph('[Figure: solver_analysis.png - Solver output]')

doc.add_paragraph(
    'The solver analysis plot shows the estimated orientation trajectory in terms of Euler angles '
    '(Yaw, Pitch, Roll). The shaded regions indicate 2-sigma uncertainty bounds derived from the '
    'filter covariance. Note that:'
)

solver_obs = [
    'Yaw (heading) has the largest uncertainty due to indoor magnetic disturbances',
    'Pitch and Roll are well-constrained by the accelerometer gravity reference',
    'Uncertainty grows during rapid motion and shrinks during stationary periods',
]
for obs in solver_obs:
    doc.add_paragraph(obs, style='List Bullet')

# Figure 5: Error Analysis
doc.add_heading('C.5 Error Distribution', level=2)
if os.path.exists('debug_output/solver_error.png'):
    doc.add_picture('debug_output/solver_error.png', width=Inches(6))
    fig5_caption = doc.add_paragraph('Figure C.5: Residual error analysis showing distribution of estimation errors.')
    fig5_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig5_caption.runs[0].font.size = Pt(10)
    fig5_caption.runs[0].font.italic = True
else:
    doc.add_paragraph('[Figure: solver_error.png - Error distribution]')

doc.add_paragraph(
    'The error analysis shows the distribution of residuals (differences between predicted and '
    'observed sensor readings). A well-tuned filter should produce residuals that are approximately '
    'Gaussian with zero mean. Systematic biases in the residuals would indicate model mismatch or '
    'uncompensated sensor errors.'
)

doc.add_page_break()

# ============================================================================
# APPENDIX D: DEVELOPMENT TIMELINE
# ============================================================================
doc.add_heading('Appendix D: Development Timeline and Challenges', level=1)

doc.add_paragraph(
    'This appendix provides a narrative account of the development process, including key '
    'milestones, challenges encountered, and lessons learned.'
)

doc.add_heading('D.1 Initial Development', level=2)

doc.add_paragraph(
    'Phase 1 (Days 1-3): The project began with implementing the PyTorch factor graph solver. '
    'Initial development focused on the basic quaternion kinematics and loss function formulation. '
    'Early tests on synthetic data showed promising results with zero-bias data, but significant '
    'problems emerged when bias was introduced.'
)

doc.add_paragraph(
    'Challenge: Weight Calculation Bug. The first major bug was an incorrect weighting of the '
    'gyroscope loss term. The original implementation used 1/σ² as the weight, but the correct '
    'formulation requires 1/(σ·dt)² to account for discrete-time integration. This single bug '
    'caused the optimizer to completely fail at bias estimation, producing either zero bias or '
    'wildly oscillating values.'
)

doc.add_paragraph(
    'Resolution: Careful rederivation of the discrete-time loss function from first principles, '
    'comparing against the continuous-time formulation in academic references [3, 4]. The fix was '
    'a one-line change but required significant debugging time to identify.'
)

doc.add_heading('D.2 Synchronization Challenge', level=2)

doc.add_paragraph(
    'Phase 2 (Days 4-5): With the solver working on synthetic data, we moved to real data from '
    'a Google Pixel 10. Immediately, a new problem emerged: the estimated orientation was clearly '
    'wrong, showing rotations that didn\'t match the video.'
)

doc.add_paragraph(
    'Challenge: Video-Sensor Time Offset. The video and sensor data had different timestamps, '
    'with an unknown offset of approximately 1-2 seconds. Multiple failed attempts to estimate '
    'the offset from file metadata led to developing the optical flow cross-correlation method.'
)

doc.add_paragraph(
    'Resolution: The sync_video.py script was developed to extract optical flow magnitude from '
    'video frames and correlate with gyroscope magnitude. This provided robust offset estimation '
    'to within ±0.1 seconds. The offset was then applied as a command-line argument to the '
    'output generation script.'
)

doc.add_heading('D.3 Google EKF Implementation', level=2)

doc.add_paragraph(
    'Phase 3 (Days 6-7): To compare against a production-quality baseline, we implemented a Python '
    'port of the Android AOSP sensor fusion algorithm. This required careful study of the C++ source '
    'code in the Android Open Source Project repository.'
)

doc.add_paragraph(
    'Challenge: Quaternion Convention Mismatch. The AOSP code uses a different quaternion '
    'multiplication convention than typical robotics references. This led to incorrect rotations '
    'until the convention was identified and matched.'
)

doc.add_paragraph(
    'Resolution: Systematic testing with known rotations (90° about each axis) to verify that '
    'the implementation matched expected behavior. Added detailed comments in the code explaining '
    'the convention choices.'
)

doc.add_heading('D.4 Noise Database Creation', level=2)

doc.add_paragraph(
    'Phase 4 (Day 8): To enable device-specific calibration, we built a comprehensive noise '
    'database covering 100+ smartphone models. This required researching which IMU chips are '
    'used in each phone and obtaining datasheet specifications for each chip.'
)

doc.add_paragraph(
    'Challenge: Proprietary Information. Phone manufacturers do not publish which IMU sensor '
    'is used in each model. We relied on a combination of teardown reports, user-posted sensor '
    'app screenshots, and brand-level patterns (e.g., Samsung flagships typically use STM sensors).'
)

doc.add_paragraph(
    'Resolution: Implemented a fallback hierarchy in the database that provides reasonable '
    'defaults even for unknown devices: exact model → partial match → brand default → generic.'
)

doc.add_heading('D.5 Final Integration', level=2)

doc.add_paragraph(
    'Phase 5 (Day 9): Final integration of all components, creation of the estimate_orientation.py '
    'simple interface, and generation of this documentation. The project now provides a complete, '
    'production-ready solution for smartphone orientation estimation.'
)

doc.add_paragraph('Lessons Learned:')

lessons = [
    'Always test on synthetic data with known ground truth before moving to real data',
    'Time synchronization between different data sources is critical and often overlooked',
    'Sensor fusion algorithms are sensitive to noise parameter tuning - good defaults matter',
    'Documentation written during development is more accurate than post-hoc documentation',
    'Iterative coordinate descent can be more stable than joint optimization for stiff problems',
]
for l in lessons:
    doc.add_paragraph(l, style='List Bullet')

# (Final save at end of script)


# ============================================================================
# SAVE FINAL DOCUMENT
# ============================================================================
doc.save('docs/CameraOrientation_Full_Documentation.docx')
fsize = os.path.getsize('docs/CameraOrientation_Full_Documentation.docx')
print(f'Complete documentation generated: docs/CameraOrientation_Full_Documentation.docx ({fsize:,} bytes)')
